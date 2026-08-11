"""보고서 최종안을 .docx 파일 형태로 출력합니다."""

import os
from docx import Document as DocxDocument

from reports.models import Report

def export_to_docx(report: Report) -> str:
    doc = DocxDocument()

    doc.add_heading(report.title or report.keyword, level=1)
    doc.add_paragraph(report.content)
    os.makedirs("exports", exist_ok=True)
    file_path = f"exports/report_{report.id}.docx"
    doc.save(file_path)

    report.file_path = file_path
    report.file_format = "docx"
    report.save()

    return file_path